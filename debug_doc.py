import os
import subprocess
import uuid
import shutil
from docx import Document

def extract_text_from_docx(path):
    print(f"Reading from: {path}")
    doc = Document(path)
    paras = []
    idx = 0
    for p in doc.paragraphs:
        if p.text:
            paras.append({"text": p.text, "index": idx})
            idx += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        paras.append({"text": p.text, "index": idx})
                        idx += 1
    return paras

def test_flow():
    # 1. Create a dummy txt file
    with open("test_input.txt", "w") as f:
        f.write("This is a test paragraph.\n\nThis is another paragraph in the doc.")
    
    # 2. Use textutil to create a .doc file (simulating user upload)
    print("Creating sample.doc...")
    subprocess.run(["textutil", "-convert", "doc", "-output", "sample.doc", "test_input.txt"], check=True)
    
    # 3. Try to convert it back to .docx and read it (the app logic)
    print("Converting sample.doc to .docx...")
    temp_id = str(uuid.uuid4())
    docx_path = f"test_out_{temp_id}.docx"
    
    try:
        # The exact command from main.py
        subprocess.run(["textutil", "-convert", "docx", "-output", docx_path, "sample.doc"], 
                      capture_output=True, text=True, check=True)
        
        paragraphs = extract_text_from_docx(docx_path)
        print(f"Extracted segments: {len(paragraphs)}")
        for p in paragraphs:
            print(f" - {p['text']}")
            
        if not paragraphs:
            print("FAILED: No segments extracted.")
        else:
            print("SUCCESS: Content extracted correctly.")
            
    except Exception as e:
        print(f"ERROR: {str(e)}")
    finally:
        # Cleanup
        for f in ["test_input.txt", "sample.doc", docx_path]:
            if os.path.exists(f): os.remove(f)

if __name__ == "__main__":
    test_flow()
