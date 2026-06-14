import os
import shutil
import uuid
import json
import asyncio
import subprocess
import traceback
import re
import pypandoc
from typing import List, Dict, Any, Optional

from fastapi import FastAPI, UploadFile, File, WebSocket, WebSocketDisconnect, HTTPException, Form
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse
from docx import Document
from docx.shared import RGBColor
from pypdf import PdfReader
from langchain_ollama import ChatOllama
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
import chromadb

app = FastAPI()

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# ChromaDB Configuration
chroma_client = chromadb.PersistentClient(path="./chroma_db")
collection = chroma_client.get_or_create_collection(name="document_reviews")

# Session State Storage
sessions: Dict[str, Dict[str, Any]] = {}

PRE_REVIEW_PROMPT = ChatPromptTemplate.from_messages([
    ("system", "你是一位资深学术期刊主编及资深新闻采编总监。请严格按照以下【专家评审准则】评估文档是否达到发表标准：\n"
               "1. 【结构逻辑】：检查是否符合 IMRaD (引言、方法、结果、讨论) 规范；逻辑是否自洽；是否存在“Nut Graf”（核心意义段落）清晰解释研究/报道的价值。\n"
               "2. 【创新与贡献】：评估研究/内容是否具有独特性，是否提供了实质性的新知识或新视角。\n"
               "3. 【严谨性】：方法论是否可复现；数据/事实是否有可靠来源支撑；结论是否过度推导。\n"
               "4. 【合规性】：检查学术伦理、利益冲突声明及引用规范。\n"
               "必须使用中文回答。格式为 JSON 对象：{{'decision': 'Pass' 或 'Fail', 'summary': '深度专家分析报告...', 'recommendations': ['具体改进点1', '具体改进点2']}}。"),
    ("user", "{text}")
])

REVIEW_PROMPT = ChatPromptTemplate.from_messages([
    ("system", "你是一位专业的稿件深度编辑。请应用以下【编辑实务规范】为段落提供修改建议：\n"
               "1. 【去冗增效】：剔除赘余词汇，修正“被动语态堆砌”，增强动词表现力，提升表达的简洁度。\n"
               "2. 【学术/专业调性】：确保语气客观严谨（学术类）或直接生动（新闻类）；消除口语化表达。\n"
               "3. 【逻辑转承】：优化句间衔接，确保论证链条环环相扣。\n"
               "4. 【事实准确性】：指出潜在的数字矛盾或陈述歧义。\n"
               "必须使用中文回答。返回 JSON 对象列表，每个对象包含：'target' (需修改的原始文本), 'suggestion' (修改后的专业版本), 'comment' (基于编辑规范的深度点评), 'start' (起始索引), 'end' (结束索引)。"
               "若无需修改则返回 []。"),
    ("user", "{text}")
])

INSTRUCTION_PROMPT = ChatPromptTemplate.from_messages([
    ("system", "你正在协助用户审阅文档。用户会就之前的建议或一般编辑规则给你指示。"
               "上下文: {context}。"
               "用户指示: {instruction}。"
               "分析指示并决定是否拒绝、接受之前的建议，或者以不同的方式修改文本。"
               "必须使用中文回答。返回 JSON 响应，包含 'action' ('reject_suggestion', 'accept_suggestion', 'modify_text', 'chat_only'), 'id' (如果涉及特定建议), 'new_text' (如果需要修改), 以及 'reply' (你给用户的中文解释)。"),
    ("user", "{instruction}")
])

def get_llm(model_name: str, api_key: Optional[str] = None):
    """Dynamically initialize LLM based on model name and provider."""
    if model_name.startswith("gemini") or model_name.startswith("gemma-7b") or model_name.startswith("gemma-2") or "gemma-4" in model_name:
        if not api_key:
            raise ValueError("Google API Key is required for cloud-hosted models")
        return ChatGoogleGenerativeAI(model=model_name, google_api_key=api_key, temperature=0)
    else:
        return ChatOllama(model=model_name, temperature=0)

def extract_text_from_docx_robust(path):
    """Aggressively extract text from paragraphs and nested tables."""
    try:
        doc = Document(path)
    except Exception as e:
        print(f"python-docx failed to open {path}: {e}")
        return []

    paras = []
    idx = 0

    def add_para(text):
        nonlocal idx
        if text and text.strip():
            paras.append({"text": text.strip(), "index": idx})
            idx += 1

    # 1. Standard paragraphs
    for p in doc.paragraphs:
        add_para(p.text)

    # 2. Deep Table extraction
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    add_para(p.text)
                for nested_table in cell.tables:
                    process_table(nested_table)

    for table in doc.tables:
        process_table(table)
        
    return paras

def split_into_paragraphs(text):
    """Splits raw text into paragraphs based on multiple newlines."""
    # Split by 2 or more newlines
    parts = re.split(r'\n\s*\n', text)
    paras = []
    for i, p in enumerate(parts):
        content = p.strip()
        if content:
            paras.append({"text": content, "index": i})
    return paras

@app.get("/")
async def get_index():
    if os.path.exists("static/index.html"):
        with open("static/index.html") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>Index file not found</h1>")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), client_id: str = Form(...)):
    try:
        allowed_exts = (".docx", ".doc", ".pdf")
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_exts:
            raise HTTPException(status_code=400, detail=f"不支持的格式: {ext}")
        
        temp_id = str(uuid.uuid4())
        temp_path = f"temp_{temp_id}{ext}"
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        paragraphs = []
        final_temp_path = temp_path
        
        if ext == ".docx":
            paragraphs = extract_text_from_docx_robust(temp_path)
            print(f"DEBUG: Found {len(paragraphs)} segments in .docx")
            
        elif ext == ".doc":
            docx_path = f"temp_{temp_id}.docx"
            try:
                # Primary attempt: pypandoc (industrial grade)
                print(f"DEBUG: Attempting pypandoc conversion for {temp_path}")
                pypandoc.convert_file(temp_path, 'docx', outputfile=docx_path)
                paragraphs = extract_text_from_docx_robust(docx_path)
                
                # Secondary fallback: textutil
                if not paragraphs:
                    print("DEBUG: pypandoc empty, trying textutil")
                    subprocess.run(["textutil", "-convert", "docx", "-output", docx_path, temp_path], 
                                 capture_output=True, text=True, check=True)
                    paragraphs = extract_text_from_docx_robust(docx_path)

                # Final fallback: plain text extraction
                if not paragraphs:
                    print("DEBUG: docx conversion failed to yield text, falling back to TXT extraction")
                    result = subprocess.run(["textutil", "-convert", "txt", "-stdout", temp_path], 
                                         capture_output=True, text=True, check=True)
                    paragraphs = split_into_paragraphs(result.stdout)
                
                if os.path.exists(temp_path): os.remove(temp_path)
                final_temp_path = docx_path
            except Exception as e:
                print(f"Conversion/Parsing error for .doc: {str(e)}")
                # Extreme fallback
                result = subprocess.run(["textutil", "-stdout", "-convert", "txt", temp_path], capture_output=True, text=True)
                paragraphs = split_into_paragraphs(result.stdout)
                final_temp_path = None
                
        elif ext == ".pdf":
            try:
                reader = PdfReader(temp_path)
                full_pdf_text = ""
                for page in reader.pages:
                    full_pdf_text += (page.extract_text() or "") + "\n\n"
                paragraphs = split_into_paragraphs(full_pdf_text)
                print(f"DEBUG: Found {len(paragraphs)} segments in .pdf")
            except Exception as e:
                print(f"PDF extraction error: {e}")
        
        if not paragraphs:
            raise Exception("无法从文件中提取任何文字内容，请确认文件是否加密、已损坏或者是受保护的 PDF。")

        if client_id not in sessions:
            sessions[client_id] = {
                "paragraphs": [],
                "suggestions": [],
                "chat_history": [],
                "temp_path": None,
                "original_text": "",
                "model_name": "llama3",
                "api_key": None
            }
        
        sessions[client_id]["temp_path"] = final_temp_path
        sessions[client_id]["paragraphs"] = paragraphs
        sessions[client_id]["original_text"] = "\n\n".join([p["text"] for p in paragraphs])
        
        return {"paragraphs": paragraphs}

    except Exception as e:
        print(f"UPLOAD ERROR: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.websocket("/ws/{client_id}")
async def websocket_endpoint(websocket: WebSocket, client_id: str):
    await websocket.accept()
    
    if client_id not in sessions:
        sessions[client_id] = {
            "paragraphs": [],
            "suggestions": [],
            "chat_history": [],
            "temp_path": None,
            "original_text": "",
            "model_name": "llama3",
            "api_key": None
        }
    
    try:
        while True:
            data = await websocket.receive_text()
            message = json.loads(data)
            
            if message["type"] == "start_review":
                sessions[client_id]["model_name"] = message.get("model", "llama3")
                sessions[client_id]["api_key"] = message.get("api_key")
                mode = message.get("mode", "detail_review")
                asyncio.create_task(run_review_process(websocket, client_id, mode))
            
            elif message["type"] == "chat":
                await handle_chat_instruction(websocket, client_id, message["text"])
                
    except WebSocketDisconnect:
        pass

async def run_review_process(websocket: WebSocket, client_id: str, mode: str):
    await websocket.send_json({"type": "status", "active": True})
    session = sessions[client_id]
    paragraphs = session["paragraphs"]
    full_text = session["original_text"]
    model_name = session["model_name"]
    api_key = session["api_key"]
    
    parser = JsonOutputParser()
    
    try:
        try:
            llm = get_llm(model_name, api_key)
        except Exception as e:
            await websocket.send_json({"type": "chat", "text": f"错误：模型初始化失败: {str(e)}"})
            return
        
        if mode == "pre_review":
            chain = PRE_REVIEW_PROMPT | llm | parser
            try:
                result = await chain.ainvoke({"text": full_text})
                await websocket.send_json({
                    "type": "pre_review_result",
                    **result
                })
            except Exception as e:
                print(f"Error in pre-review: {e}")
                await websocket.send_json({"type": "chat", "text": f"初审分析出错: {str(e)}"})
        else:
            chain = REVIEW_PROMPT | llm | parser
            for i, para in enumerate(paragraphs):
                try:
                    result = await chain.ainvoke({"text": para["text"]})
                    if isinstance(result, list):
                        for sug in result:
                            sug_id = str(uuid.uuid4())
                            sug["id"] = sug_id
                            sug["para_index"] = i
                            sug["status"] = "pending"
                            session["suggestions"].append(sug)
                            await websocket.send_json({
                                "type": "suggestion",
                                **sug
                            })
                except Exception as e:
                    print(f"Error reviewing paragraph {i}: {e}")
                    continue
    finally:
        await websocket.send_json({"type": "status", "active": False})

async def handle_chat_instruction(websocket: WebSocket, client_id: str, instruction: str):
    await websocket.send_json({"type": "status", "active": True})
    session = sessions[client_id]
    session["chat_history"].append({"role": "user", "content": instruction})
    
    model_name = session["model_name"]
    api_key = session["api_key"]
    
    try:
        try:
            llm = get_llm(model_name, api_key)
        except Exception as e:
            await websocket.send_json({"type": "chat", "text": f"错误：模型初始化失败: {str(e)}"})
            return
        
        context = {
            "recent_suggestions": session["suggestions"][-10:],
            "paragraphs": [p["text"] for p in session["paragraphs"][:5]]
        }
        
        parser = JsonOutputParser()
        chain = INSTRUCTION_PROMPT | llm | parser
        
        try:
            response = await chain.ainvoke({
                "context": json.dumps(context),
                "instruction": instruction
            })
            
            reply = response.get("reply", "我已处理您的指令。")
            action = response.get("action")
            sug_id = response.get("id")
            
            if action == "reject_suggestion" and sug_id:
                for s in session["suggestions"]:
                    if s["id"] == sug_id:
                        s["status"] = "rejected"
                        await notify_update(websocket, client_id, s["para_index"])
                        break
            elif action == "accept_suggestion" and sug_id:
                for s in session["suggestions"]:
                    if s["id"] == sug_id:
                        s["status"] = "accepted"
                        await notify_update(websocket, client_id, s["para_index"])
                        break
                        
            await websocket.send_json({"type": "chat", "text": reply})
            session["chat_history"].append({"role": "ai", "content": reply})
            
        except Exception as e:
            print(f"Error handling chat: {e}")
            error_msg = f"抱歉，处理指令时出现问题: {str(e)}"
            await websocket.send_json({"type": "chat", "text": error_msg})
    finally:
        await websocket.send_json({"type": "status", "active": False})

async def notify_update(websocket: WebSocket, client_id: str, para_index: int):
    para = sessions[client_id]["paragraphs"][para_index]
    para_suggestions = [s for s in sessions[client_id]["suggestions"] if s["para_index"] == para_index]
    await websocket.send_json({
        "type": "update",
        "index": para_index,
        "text": para["text"],
        "suggestions": para_suggestions
    })

@app.post("/export/{client_id}")
async def export_document(client_id: str):
    if client_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = sessions[client_id]
    
    final_doc = Document()
    for i, para_data in enumerate(session["paragraphs"]):
        p = final_doc.add_paragraph()
        text = para_data["text"]
        
        sugs = [s for s in session["suggestions"] if s["para_index"] == i and s["status"] == "accepted"]
        sugs.sort(key=lambda x: x["start"])
        
        last_idx = 0
        for s in sugs:
            p.add_run(text[last_idx:s["start"]])
            run_orig = p.add_run(text[s["start"]:s["end"]])
            run_orig.font.strike = True
            run_orig.font.color.rgb = RGBColor(128, 128, 128)
            
            run_sug = p.add_run(f" [{s['suggestion']}] ")
            run_sug.font.color.rgb = RGBColor(0, 102, 204)
            run_sug.font.bold = True
            
            last_idx = s["end"]
        p.add_run(text[last_idx:])

    export_path = f"reviewed_{client_id}.docx"
    final_doc.save(export_path)
    
    await save_to_vector_db(client_id)
    
    return FileResponse(export_path, filename="reviewed_document.docx")

async def save_to_vector_db(client_id: str):
    session = sessions[client_id]
    content = {
        "original_document": session["original_text"],
        "final_suggestions": [s for s in session["suggestions"] if s["status"] == "accepted"],
        "chat_history": session["chat_history"]
    }
    
    collection.add(
        documents=[json.dumps(content)],
        metadatas=[{"client_id": client_id, "timestamp": str(uuid.uuid4())}],
        ids=[client_id]
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
