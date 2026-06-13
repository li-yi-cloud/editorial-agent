import os
import pytest
import json
from unittest.mock import MagicMock, patch
from docx import Document
from fastapi.testclient import TestClient
from main import app, sessions, get_llm

client = TestClient(app)

@pytest.fixture
def sample_docx(tmp_path):
    file_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is the first paragraph.")
    doc.save(file_path)
    return file_path

def test_read_main():
    response = client.get("/")
    assert response.status_code == 200
    assert "AI Document Reviewer" in response.text

def test_upload_docx(sample_docx):
    with open(sample_docx, "rb") as f:
        response = client.post(
            "/upload",
            files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"client_id": "test_user"}
        )
    assert response.status_code == 200
    data = response.json()
    assert "paragraphs" in data
    assert len(data["paragraphs"]) == 1

def test_get_llm_logic():
    # Test local model selection
    llm_local = get_llm("llama3")
    from langchain_ollama import ChatOllama
    assert isinstance(llm_local, ChatOllama)
    
    # Test cloud model selection (should raise error without key)
    with pytest.raises(ValueError, match="Google API Key is required"):
        get_llm("gemini-1.5-pro")
    
    # Test cloud model selection with key (using mock since we don't want real API calls)
    with patch('main.ChatGoogleGenerativeAI') as mock_google:
        get_llm("gemini-1.5-pro", api_key="fake_key")
        mock_google.assert_called_once()
        
    with patch('main.ChatGoogleGenerativeAI') as mock_google:
        get_llm("gemma-4-31b", api_key="fake_key")
        mock_google.assert_called_once()

@pytest.mark.asyncio
async def test_export_document_format(tmp_path):
    client_id = "test_export_fmt"
    sessions[client_id] = {
        "paragraphs": [{"text": "Original.", "index": 0}],
        "suggestions": [{
            "id": "sug1",
            "para_index": 0,
            "target": "Original.",
            "suggestion": "Better.",
            "comment": "C",
            "start": 0,
            "end": 8,
            "status": "accepted"
        }],
        "chat_history": [],
        "temp_path": None,
        "original_text": "Original.",
        "model_name": "llama3",
        "api_key": None
    }
    
    response = client.post(f"/export/{client_id}")
    assert response.status_code == 200
    
    # Verify file content
    export_path = tmp_path / "fmt_test.docx"
    with open(export_path, "wb") as f:
        f.write(response.content)
    
    doc = Document(export_path)
    # The first paragraph should have multiple runs (strikethrough + suggestion)
    assert len(doc.paragraphs[0].runs) >= 2
    
    if os.path.exists(f"reviewed_{client_id}.docx"):
        os.remove(f"reviewed_{client_id}.docx")

def test_upload_isolation():
    # Upload from user A
    from io import BytesIO
    doc = Document()
    doc.add_paragraph("User A Content")
    f_a = BytesIO()
    doc.save(f_a)
    f_a.seek(0)
    
    client.post("/upload", files={"file": ("a.docx", f_a, "application/vnd.docx")}, data={"client_id": "user_a"})
    
    # Upload from user B
    doc = Document()
    doc.add_paragraph("User B Content")
    f_b = BytesIO()
    doc.save(f_b)
    f_b.seek(0)
    
    client.post("/upload", files={"file": ("b.docx", f_b, "application/vnd.docx")}, data={"client_id": "user_b"})
    
    assert "user_a" in sessions
    assert "user_b" in sessions
    assert sessions["user_a"]["temp_path"] != sessions["user_b"]["temp_path"]
    
    # Clean up
    for s in sessions.values():
        if s["temp_path"] and os.path.exists(s["temp_path"]):
            os.remove(s["temp_path"])
